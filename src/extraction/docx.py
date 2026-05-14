"""DOCX extractor using python-docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore[import]

from .base import Extractor, ExtractionResult
from .pdf import _dictionary_match_score, _layout_score

__all__ = ["DOCXExtractor"]


class DOCXExtractor(Extractor):
    def extract(self, path: Path) -> ExtractionResult:
        doc = Document(str(path))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_parts:
                    parts.append(" | ".join(row_parts))

        text = "\n".join(parts)

        if not text.strip():
            return ExtractionResult(
                text="",
                quality=0.0,
                metadata={"extractor": "python-docx"},
            )

        q_dict = _dictionary_match_score(text)
        q_layout = _layout_score(text)
        quality = round(q_dict * q_layout, 4)
        quality = max(quality, 0.1)

        return ExtractionResult(
            text=text,
            quality=quality,
            metadata={"extractor": "python-docx"},
        )
